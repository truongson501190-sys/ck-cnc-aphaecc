from flask import Flask, request, jsonify
from flask_cors import CORS
import easyocr
import tempfile
import os
import traceback
import re
from pathlib import Path

app = Flask(__name__)
CORS(app)

print("⏳ Đang tải EasyOCR...")
reader = easyocr.Reader(['vi', 'en'], gpu=False, verbose=False)
print("✅ EasyOCR sẵn sàng!")

# ============================================
# KIỂM TRA CÁC THƯ VIỆN HỖ TRỢ
# ============================================
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
    print("✅ PyMuPDF có sẵn - hỗ trợ PDF")
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("⚠️ PyMuPDF chưa cài - không hỗ trợ PDF (cài: pip install PyMuPDF)")

try:
    from docx import Document
    DOCX_AVAILABLE = True
    print("✅ python-docx có sẵn - hỗ trợ DOCX")
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx chưa cài - không hỗ trợ DOCX (cài: pip install python-docx)")

try:
    import openpyxl
    EXCEL_AVAILABLE = True
    print("✅ openpyxl có sẵn - hỗ trợ XLSX")
except ImportError:
    EXCEL_AVAILABLE = False
    print("⚠️ openpyxl chưa cài - không hỗ trợ XLSX (cài: pip install openpyxl)")

# ============================================
# HÀM XỬ LÝ CÁC LOẠI FILE
# ============================================

def extract_text_from_image(image_path):
    """OCR từ ảnh bằng EasyOCR"""
    try:
        result = reader.readtext(image_path, detail=0)
        text = '\n'.join(result)
        return text.strip()
    except Exception as e:
        print(f"❌ Lỗi OCR ảnh: {e}")
        return None

def extract_text_from_pdf(pdf_path):
    """Trích xuất văn bản từ PDF (hỗ trợ cả scan)"""
    if not PYMUPDF_AVAILABLE:
        return None
    
    try:
        doc = fitz.open(pdf_path)
        all_text = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Thử lấy text trực tiếp (nếu PDF có text layer)
            text = page.get_text()
            if text.strip():
                all_text.append(f"--- Trang {page_num+1} ---\n{text.strip()}")
            else:
                # Nếu không có text, dùng OCR
                pix = page.get_pixmap(dpi=200)
                img_data = pix.tobytes("png")
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                    tmp.write(img_data)
                    tmp_path = tmp.name
                
                ocr_text = extract_text_from_image(tmp_path)
                os.unlink(tmp_path)
                if ocr_text:
                    all_text.append(f"--- Trang {page_num+1} (OCR) ---\n{ocr_text}")
        
        doc.close()
        return '\n\n'.join(all_text)
    except Exception as e:
        print(f"❌ Lỗi xử lý PDF: {e}")
        return None

def extract_text_from_docx(docx_path):
    """Trích xuất văn bản từ file DOCX"""
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document(docx_path)
        text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        # Đọc cả bảng nếu có
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text for cell in row.cells if cell.text.strip()])
                if row_text:
                    text += '\n' + row_text
        return text.strip()
    except Exception as e:
        print(f"❌ Lỗi xử lý DOCX: {e}")
        return None

def extract_text_from_excel(excel_path):
    """Trích xuất văn bản từ file Excel (XLSX)"""
    if not EXCEL_AVAILABLE:
        return None
    
    try:
        wb = openpyxl.load_workbook(excel_path, data_only=True)
        all_text = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_text = [f"--- Sheet: {sheet_name} ---"]
            
            for row in sheet.iter_rows(values_only=True):
                row_text = ' | '.join([str(cell) for cell in row if cell is not None])
                if row_text:
                    sheet_text.append(row_text)
            
            if len(sheet_text) > 1:
                all_text.append('\n'.join(sheet_text))
        
        wb.close()
        return '\n\n'.join(all_text)
    except Exception as e:
        print(f"❌ Lỗi xử lý Excel: {e}")
        return None

def extract_text_from_txt(txt_path):
    """Đọc file TXT đơn giản"""
    try:
        with open(txt_path, 'r', encoding='utf-8') as f:
            return f.read()
    except:
        try:
            with open(txt_path, 'r', encoding='latin-1') as f:
                return f.read()
        except:
            return None

# ============================================
# HÀM PHÂN LOẠI VÀ XỬ LÝ FILE
# ============================================

def process_file(filepath, filename):
    """Xử lý file dựa trên phần mở rộng"""
    ext = filename.lower().split('.')[-1]
    
    # Ảnh
    if ext in ['jpg', 'jpeg', 'png', 'bmp', 'tiff', 'tif', 'webp', 'heic', 'heif', 'gif']:
        return extract_text_from_image(filepath), 'image'
    
    # PDF
    elif ext == 'pdf':
        return extract_text_from_pdf(filepath), 'pdf'
    
    # Word
    elif ext in ['docx', 'doc']:
        return extract_text_from_docx(filepath), 'document'
    
    # Excel
    elif ext in ['xlsx', 'xls']:
        return extract_text_from_excel(filepath), 'spreadsheet'
    
    # Text
    elif ext in ['txt', 'csv', 'md', 'log']:
        return extract_text_from_txt(filepath), 'text'
    
    else:
        return None, 'unknown'

# ============================================
# API ENDPOINTS
# ============================================

@app.route('/health', methods=['GET'])
def health():
    return jsonify({
        'status': 'ok',
        'supported_formats': {
            'image': ['jpg', 'jpeg', 'png', 'bmp', 'tiff', 'webp', 'heic'],
            'pdf': ['pdf'],
            'document': ['docx'],
            'spreadsheet': ['xlsx'],
            'text': ['txt', 'csv']
        }
    })

@app.route('/ocr', methods=['POST'])
def ocr():
    try:
        print("📥 Nhận request OCR")
        
        if 'file' not in request.files:
            print("❌ Không có file trong request")
            return jsonify({'status': 'error', 'message': 'No file'}), 400
        
        file = request.files['file']
        print(f"📎 File nhận: {file.filename}, type: {file.content_type}, size: {file.content_length}")
        
        if file.filename == '':
            print("❌ Tên file rỗng")
            return jsonify({'status': 'error', 'message': 'Empty filename'}), 400

        # Lưu file tạm (giữ nguyên đuôi mở rộng)
        original_ext = os.path.splitext(file.filename)[1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=original_ext) as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name
            print(f"💾 Lưu file tạm: {tmp_path}")

        # Xử lý file theo loại
        text, file_type = process_file(tmp_path, file.filename)
        
        # Xóa file tạm
        os.unlink(tmp_path)
        print("🗑️ Xóa file tạm")
        
        if text is None:
            return jsonify({
                'status': 'error',
                'message': f'Không thể xử lý file .{original_ext}',
                'supported_types': ['jpg', 'jpeg', 'png', 'bmp', 'tiff', 'pdf', 'docx', 'xlsx', 'txt']
            }), 400

        # Làm sạch text
        text = re.sub(r'\s+', ' ', text).strip()

        return jsonify({
            'status': 'success' if text else 'error',
            'text': text or 'Không tìm thấy nội dung',
            'length': len(text),
            'filename': file.filename,
            'file_type': file_type,
            'lines': len(text.split('\n')) if text else 0
        })

    except Exception as e:
        print("❌ LỖI NỘI BỘ:")
        traceback.print_exc()
        return jsonify({'status': 'error', 'message': str(e)}), 500

if __name__ == '__main__':
    print("=" * 60)
    print("🚀 OCR Server Advanced")
    print("📌 Health: GET /health")
    print("📌 OCR: POST /ocr")
    print("📌 Hỗ trợ: ảnh, PDF, DOCX, XLSX, TXT")
    print("=" * 60)
    app.run(host='0.0.0.0', port=5001, debug=False)