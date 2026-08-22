from __future__ import annotations

import os
import re
import tempfile
import traceback
from typing import Optional, Tuple

from flask import Flask, jsonify, request
from flask_cors import CORS

try:
    import easyocr
except ImportError:  # pragma: no cover
    easyocr = None

try:
    import fitz
except ImportError:  # pragma: no cover
    fitz = None

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None

try:
    import openpyxl
except ImportError:  # pragma: no cover
    openpyxl = None

app = Flask(__name__)
CORS(app)
app.config["MAX_CONTENT_LENGTH"] = 100 * 1024 * 1024

reader = easyocr.Reader(["vi", "en"], gpu=False, verbose=False) if easyocr else None

PYMUPDF_AVAILABLE = fitz is not None
DOCX_AVAILABLE = Document is not None
EXCEL_AVAILABLE = openpyxl is not None


def extract_text_from_image(image_path: str) -> Optional[str]:
    """OCR image content using EasyOCR."""
    if reader is None:
        return None
    try:
        result = reader.readtext(image_path, detail=0)
        text = "\n".join(result)
        return text.strip()
    except Exception as exc:  # pragma: no cover
        print(f"❌ Lỗi OCR ảnh: {exc}")
        return None


def extract_text_from_pdf(pdf_path: str) -> Optional[str]:
    """Extract text from PDF, converting pages to images if needed."""
    if not PYMUPDF_AVAILABLE:
        return None
    try:
        doc = fitz.open(pdf_path)
        all_text = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                all_text.append(f"--- Trang {page_num + 1} ---\n{text.strip()}")
                continue

            pix = page.get_pixmap(dpi=200)
            img_data = pix.tobytes("png")
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp.write(img_data)
                tmp_path = tmp.name

            ocr_text = extract_text_from_image(tmp_path)
            os.unlink(tmp_path)
            if ocr_text:
                all_text.append(f"--- Trang {page_num + 1} (OCR) ---\n{ocr_text}")

        doc.close()
        return "\n\n".join(all_text)
    except Exception as exc:  # pragma: no cover
        print(f"❌ Lỗi xử lý PDF: {exc}")
        traceback.print_exc()
        return None


def extract_text_from_docx(docx_path: str) -> Optional[str]:
    """Extract text from DOCX file."""
    if not DOCX_AVAILABLE:
        return None
    try:
        doc = Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells if cell.text.strip()])
                if row_text:
                    text += "\n" + row_text
        return text.strip()
    except Exception as exc:  # pragma: no cover
        print(f"❌ Lỗi xử lý DOCX: {exc}")
        return None


def extract_text_from_excel(excel_path: str) -> Optional[str]:
    """Extract text from XLSX file."""
    if not EXCEL_AVAILABLE:
        return None
    try:
        wb = openpyxl.load_workbook(excel_path, data_only=True)
        all_text = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_text = [f"--- Sheet: {sheet_name} ---"]
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join([str(cell) for cell in row if cell is not None])
                if row_text:
                    sheet_text.append(row_text)
            if len(sheet_text) > 1:
                all_text.append("\n".join(sheet_text))
        wb.close()
        return "\n\n".join(all_text)
    except Exception as exc:  # pragma: no cover
        print(f"❌ Lỗi xử lý Excel: {exc}")
        return None


def extract_text_from_txt(txt_path: str) -> Optional[str]:
    """Read text from TXT file."""
    try:
        with open(txt_path, "r", encoding="utf-8") as file:
            return file.read()
    except UnicodeDecodeError:
        try:
            with open(txt_path, "r", encoding="latin-1") as file:
                return file.read()
        except Exception:
            return None
    except Exception:
        return None


def process_file(filepath: str, filename: str) -> Tuple[Optional[str], str]:
    """Classify a file by extension and extract its text."""
    ext = filename.lower().split(".")[-1]

    if ext in ["jpg", "jpeg", "png", "bmp", "tiff", "tif", "webp", "heic", "heif", "gif"]:
        return extract_text_from_image(filepath), "image"
    if ext == "pdf":
        return extract_text_from_pdf(filepath), "pdf"
    if ext in ["docx", "doc"]:
        return extract_text_from_docx(filepath), "document"
    if ext in ["xlsx", "xls"]:
        return extract_text_from_excel(filepath), "spreadsheet"
    if ext in ["txt", "csv", "md", "log"]:
        return extract_text_from_txt(filepath), "text"
    return None, "unknown"


@app.route("/health", methods=["GET"])
def health():
    return jsonify({
        "status": "ok",
        "easyocr": reader is not None,
        "pdf_support": PYMUPDF_AVAILABLE,
        "docx_support": DOCX_AVAILABLE,
        "excel_support": EXCEL_AVAILABLE,
        "max_size_mb": 100,
    })


@app.route("/ocr", methods=["POST", "OPTIONS"])
def ocr():
    if request.method == "OPTIONS":
        return "", 200

    if "file" not in request.files:
        return jsonify({"status": "error", "message": "No file uploaded"}), 400

    uploaded_file = request.files["file"]
    if uploaded_file.filename == "":
        return jsonify({"status": "error", "message": "Empty filename"}), 400

    suffix = os.path.splitext(uploaded_file.filename)[1] or ".jpg"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        uploaded_file.save(tmp.name)
        tmp_path = tmp.name

    try:
        text, _ = process_file(tmp_path, uploaded_file.filename)
        if text is None:
            return jsonify({"status": "error", "message": "Không thể trích xuất văn bản"}), 500

        cleaned = re.sub(r"\s+", " ", text).strip()
        return jsonify({
            "status": "success" if cleaned else "error",
            "text": cleaned or "⚠️ Không tìm thấy chữ",
            "length": len(cleaned),
            "filename": uploaded_file.filename,
        })
    except Exception as exc:  # pragma: no cover
        print(f"❌ Lỗi OCR server: {exc}")
        return jsonify({"status": "error", "message": str(exc)}), 500
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5001, debug=False)
