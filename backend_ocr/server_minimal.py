# backend_ocr/server_minimal.py
from flask import Flask, request, jsonify
from flask_cors import CORS
import easyocr
import tempfile
import os
import traceback

app = Flask(__name__)
CORS(app)

# ✅ Tăng giới hạn kích thước file lên 100MB
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB

print("⏳ Đang tải EasyOCR...")
reader = easyocr.Reader(['vi', 'en'], gpu=False, verbose=False)
print("✅ EasyOCR sẵn sàng!")

# Kiểm tra PyMuPDF
try:
    import fitz
    PYMUPDF_AVAILABLE = True
    print("✅ PyMuPDF có sẵn - hỗ trợ PDF scan")
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("⚠️ PyMuPDF chưa cài - không hỗ trợ PDF scan (cài: pip install PyMuPDF)")

# Kiểm tra python-docx (hỗ trợ Word)
try:
    from docx import Document
    DOCX_AVAILABLE = True
    print("✅ python-docx có sẵn - hỗ trợ DOCX")
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx chưa cài - không hỗ trợ DOCX (cài: pip install python-docx)")

# Kiểm tra openpyxl (hỗ trợ Excel)
try:
    import openpyxl
    EXCEL_AVAILABLE = True
    print("✅ openpyxl có sẵn - hỗ trợ XLSX")
except ImportError:
    EXCEL_AVAILABLE = False
    print("⚠️ openpyxl chưa cài - không hỗ trợ XLSX (cài: pip install openpyxl)")

def process_pdf(pdf_path):
    """Xử lý PDF scan: chuyển sang ảnh và OCR từng trang"""
    if not PYMUPDF_AVAILABLE:
        return None
    
    try:
        doc = fitz.open(pdf_path)
        all_text = []
        total_pages = len(doc)
        
        for page_num in range(total_pages):
            print(f"📄 Đang xử lý trang {page_num + 1}/{total_pages}...")
            
            # Chuyển trang thành ảnh
            page = doc[page_num]
            pix = page.get_pixmap(dpi=200)
            img_data = pix.tobytes("png")
            
            # Lưu ảnh tạm
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                tmp.write(img_data)
                tmp_path = tmp.name
            
            # OCR trang
            result = reader.readtext(tmp_path, detail=0, paragraph=True)
            text = '\n'.join(result)
            
            # Xóa file tạm
            os.unlink(tmp_path)
            
            if text.strip():
                all_text.append(f"--- Trang {page_num + 1} ---\n{text}")
        
        doc.close()
        return '\n\n'.join(all_text)
        
    except Exception as e:
        print(f"❌ Lỗi xử lý PDF: {e}")
        traceback.print_exc()
        return None

def process_docx(filepath):
    """Xử lý file Word (DOCX)"""
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document(filepath)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        
        # Đọc cả bảng nếu có
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text for cell in row.cells if cell.text.strip()])
                if row_text:
                    text += '\n' + row_text
        
        return text.strip()
    except Exception as e:
        print(f"❌ Lỗi xử lý DOCX: {e}")
        return None

def process_excel(filepath):
    """Xử lý file Excel (XLSX)"""
    if not EXCEL_AVAILABLE:
        return None
    
    try:
        wb = openpyxl.load_workbook(filepath, data_only=True)
        all_text = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_text = [f"--- Sheet: {sheet_name} ---"]
            
            for row in sheet.iter_rows(values_only=True):
                row_text = ' | '.join([str(cell) for cell in row if cell is not None])
                if row_text:
                    sheet_text.append(row_text)
            
            if len(sheet_text) > 1:
                all_text.append('\n'.join(sheet_text))
        
        wb.close()
        return '\n\n'.join(all_text)
    except Exception as e:
        print(f"❌ Lỗi xử lý Excel: {e}")
        return None

@app.route('/health', methods=['GET'])
def health():
    return jsonify({
        'status': 'ok',
        'pdf_support': PYMUPDF_AVAILABLE,
        'docx_support': DOCX_AVAILABLE,
        'excel_support': EXCEL_AVAILABLE,
        'max_size_mb': 100
    })

@app.route('/ocr', methods=['POST', 'OPTIONS'])
def ocr():
    if request.method == 'OPTIONS':
        return '', 200
    
    tmp_path = None
    try:
        print("=" * 50)
        print("📥 Nhận request OCR")
        
        if 'file' not in request.files:
            return jsonify({'status': 'error', 'message': 'No file'}), 400
        
        file = request.files['file']
        file_size = request.content_length or 0
        print(f"📎 File: {file.filename}, type: {file.content_type}, size: {file_size} bytes ({file_size/1024/1024:.2f} MB)")
        
        if file.filename == '':
            return jsonify({'status': 'error', 'message': 'Empty filename'}), 400

        # Lưu file tạm
        suffix = os.path.splitext(file.filename)[1] or '.jpg'
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name
            print(f"💾 Lưu tạm: {tmp_path}")

        # Xác định loại file
        filename_lower = file.filename.lower()
        is_pdf = filename_lower.endswith('.pdf')
        is_docx = filename_lower.endswith('.docx') or filename_lower.endswith('.doc')
        is_excel = filename_lower.endswith('.xlsx') or filename_lower.endswith('.xls')
        is_image = filename_lower.endswith(('.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.webp', '.heic'))
        is_txt = filename_lower.endswith('.txt')
        
        text = None
        
        # Xử lý theo loại file
        if is_pdf:
            if PYMUPDF_AVAILABLE:
                print("📄 Xử lý PDF scan...")
                text = process_pdf(tmp_path)
            else:
                return jsonify({
                    'status': 'error',
                    'message': 'PyMuPDF chưa được cài đặt. Chạy: pip install PyMuPDF'
                }), 500
                
        elif is_docx:
            if DOCX_AVAILABLE:
                print("📄 Xử lý DOCX...")
                text = process_docx(tmp_path)
            else:
                return jsonify({
                    'status': 'error',
                    'message': 'python-docx chưa được cài đặt. Chạy: pip install python-docx'
                }), 500
                
        elif is_excel:
            if EXCEL_AVAILABLE:
                print("📊 Xử lý Excel...")
                text = process_excel(tmp_path)
            else:
                return jsonify({
                    'status': 'error',
                    'message': 'openpyxl chưa được cài đặt. Chạy: pip install openpyxl'
                }), 500
                
        elif is_txt:
            print("📄 Xử lý TXT...")
            try:
                with open(tmp_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            except:
                try:
                    with open(tmp_path, 'r', encoding='latin-1') as f:
                        text = f.read()
                except Exception as e:
                    print(f"❌ Lỗi đọc TXT: {e}")
                    text = None
                    
        elif is_image:
            # OCR ảnh
            print("🔍 Đang OCR ảnh...")
            result = reader.readtext(tmp_path, detail=0, paragraph=True)
            text = '\n'.join(result)
        else:
            # Không xác định, thử OCR như ảnh
            print("🔍 Không xác định định dạng, thử OCR...")
            result = reader.readtext(tmp_path, detail=0, paragraph=True)
            text = '\n'.join(result)
        
        # Xóa file tạm
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
            print("🗑️ Xóa file tạm")

        if text is None:
            return jsonify({
                'status': 'error',
                'message': 'Không thể xử lý file'
            }), 500

        print(f"✅ OCR thành công: {len(text)} ký tự")
        print(f"📝 Preview: {text[:200]}...")

        return jsonify({
            'status': 'success' if text else 'error',
            'text': text or 'Không tìm thấy chữ',
            'length': len(text),
            'filename': file.filename,
            'is_pdf': is_pdf,
            'file_type': 'pdf' if is_pdf else 'docx' if is_docx else 'excel' if is_excel else 'txt' if is_txt else 'image'
        })

    except Exception as e:
        print("❌ LỖI:")
        traceback.print_exc()
        
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except:
                pass
        
        return jsonify({'status': 'error', 'message': str(e)}), 500

if __name__ == '__main__':
    print("=" * 60)
    print("🚀 OCR Server (hỗ trợ: ảnh, PDF, DOCX, XLSX, TXT)")
    print("📌 Health: GET /health")
    print("📌 OCR: POST /ocr")
    print("📌 Giới hạn file: 100MB")
    print("=" * 60)
    app.run(host='0.0.0.0', port=5001, debug=True)